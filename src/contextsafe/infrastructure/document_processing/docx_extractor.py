"""
DOCX text extractor.

Uses python-docx for text extraction from Word documents.

Traceability:
- Contract: CNT-T3-DOCX-EXTRACTOR-001
- Port: ports.TextExtractor
"""

from __future__ import annotations

import io

from contextsafe.application.ports import ExtractionResult, TextExtractor


class DocxExtractor(TextExtractor):
    """
    DOCX text extractor using python-docx.

    Extracts:
    - Paragraphs
    - Tables
    - Headers/footers
    """

    async def extract(
        self,
        content: bytes,
        filename: str,
        ocr_fallback: bool = True,
    ) -> ExtractionResult:
        """
        Extract text from DOCX content.

        Args:
            content: Raw DOCX bytes
            filename: Original filename
            ocr_fallback: Ignored (not applicable for DOCX)

        Returns:
            ExtractionResult with extracted text
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            texts: list[str] = []
            has_tables = False
            has_images = False
            metadata: dict[str, str] = {}

            # Extract core properties (metadata)
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.created:
                    metadata["created"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified"] = core_props.modified.isoformat()
            except Exception:
                pass

            # Extract headers from first section
            try:
                for section in doc.sections:
                    header = section.header
                    if header:
                        header_text = "\n".join(
                            p.text for p in header.paragraphs if p.text.strip()
                        )
                        if header_text:
                            texts.append(f"[ENCABEZADO]\n{header_text}")
                    break  # Only first section's header
            except Exception:
                pass

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

            # Extract tables
            for table in doc.tables:
                has_tables = True
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        texts.append(row_text)

            # Extract footers from first section
            try:
                for section in doc.sections:
                    footer = section.footer
                    if footer:
                        footer_text = "\n".join(
                            p.text for p in footer.paragraphs if p.text.strip()
                        )
                        if footer_text:
                            texts.append(f"\n[PIE DE PÁGINA]\n{footer_text}")
                    break  # Only first section's footer
            except Exception:
                pass

            # Check for images (simplified check)
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref.lower():
                        has_images = True
                        break
            except Exception:
                has_images = False

            full_text = "\n\n".join(texts)

            # Estimate page count (roughly 500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)

            return ExtractionResult(
                text=full_text,
                format_detected="docx",
                page_count=page_count,
                has_tables=has_tables,
                has_images=has_images,
                ocr_used=False,
                confidence=0.95,
                metadata=metadata if metadata else None,
            )

        except ImportError:
            return ExtractionResult(
                text="",
                format_detected="docx",
                confidence=0.0,
                metadata={"error": "python-docx not installed"},
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                format_detected="docx",
                confidence=0.0,
                metadata={"error": str(e)},
            )

    async def extract_from_path(
        self,
        path: str,
        ocr_fallback: bool = True,
    ) -> ExtractionResult:
        """
        Extract text from a DOCX file.

        Args:
            path: Path to the DOCX file
            ocr_fallback: Ignored (not applicable for DOCX)

        Returns:
            ExtractionResult with extracted text
        """
        try:
            with open(path, "rb") as f:
                content = f.read()
            return await self.extract(content, path, ocr_fallback)
        except FileNotFoundError:
            return ExtractionResult(
                text="",
                format_detected="docx",
                confidence=0.0,
                metadata={"error": f"File not found: {path}"},
            )

    def supports_format(self, extension: str) -> bool:
        """Check if this extractor supports the format."""
        return extension.lower() in {".docx", ".doc"}
