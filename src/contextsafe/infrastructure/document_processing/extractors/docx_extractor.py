"""
DOCX document extractor using python-docx.

Traceability:
- IMPLEMENTATION_PLAN.md Phase 4.2
"""

from __future__ import annotations

from typing import BinaryIO

from contextsafe.infrastructure.document_processing.extractors.base import (
    DocumentExtractor,
    ExtractionResult,
)


class DocxExtractor(DocumentExtractor):
    """
    Extract text from DOCX documents.

    Uses python-docx to extract text, tables, headers, and footers.
    Preserves document structure.
    """

    @property
    def supported_extensions(self) -> list[str]:
        return ["docx", "doc"]

    @property
    def format_name(self) -> str:
        return "Word Document"

    def extract(self, file: BinaryIO, filename: str = "") -> ExtractionResult:
        """
        Extract text from DOCX document.

        Extracts paragraphs, tables, headers, and footers.
        """
        errors: list[str] = []
        text_parts: list[str] = []
        tables: list[list[list[str]]] = []
        metadata: dict[str, str] = {}

        try:
            from docx import Document

            doc = Document(file)

            # Extract core properties (metadata)
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.created:
                    metadata["created"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified"] = core_props.modified.isoformat()
            except Exception as e:
                errors.append(f"Error extracting metadata: {str(e)}")

            # Extract headers from first section
            try:
                for section in doc.sections:
                    header = section.header
                    if header:
                        header_text = "\n".join(
                            p.text for p in header.paragraphs if p.text.strip()
                        )
                        if header_text:
                            text_parts.append(f"[ENCABEZADO]\n{header_text}")
                    break  # Only first section's header
            except Exception as e:
                errors.append(f"Error extracting header: {str(e)}")

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                try:
                    table_data: list[list[str]] = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    if table_data:
                        tables.append(table_data)
                        # Also add table as text
                        table_text = self._table_to_text(table_data)
                        text_parts.append(f"\n[TABLA]\n{table_text}")
                except Exception as e:
                    errors.append(f"Error extracting table: {str(e)}")

            # Extract footers from first section
            try:
                for section in doc.sections:
                    footer = section.footer
                    if footer:
                        footer_text = "\n".join(
                            p.text for p in footer.paragraphs if p.text.strip()
                        )
                        if footer_text:
                            text_parts.append(f"\n[PIE DE PÁGINA]\n{footer_text}")
                    break
            except Exception as e:
                errors.append(f"Error extracting footer: {str(e)}")

            text = "\n\n".join(text_parts)

            # Estimate page count (rough: ~3000 chars per page)
            page_count = max(1, len(text) // 3000 + 1)

        except Exception as e:
            errors.append(f"DOCX extraction error: {str(e)}")
            text = ""
            page_count = 0

        return ExtractionResult(
            text=text,
            page_count=page_count,
            tables=tables,
            metadata=metadata,
            errors=errors,
            format="docx",
        )

    def _table_to_text(self, table: list[list[str]]) -> str:
        """Convert table to readable text format."""
        if not table:
            return ""

        # Calculate column widths
        col_widths = []
        for col_idx in range(len(table[0])):
            max_width = 0
            for row in table:
                if col_idx < len(row):
                    max_width = max(max_width, len(row[col_idx]))
            col_widths.append(min(max_width, 30))  # Cap at 30 chars

        lines = []
        for row in table:
            cells = []
            for i, cell in enumerate(row):
                width = col_widths[i] if i < len(col_widths) else 20
                cells.append(cell[:width].ljust(width))
            lines.append(" | ".join(cells))

        return "\n".join(lines)
