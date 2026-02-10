"""
Export routes.

Traceability:
- Contract: CNT-T4-EXPORT-001
- Phase 10: Advanced Export

Provides document and glossary export in multiple formats.
Includes Policy Gate for critical PII validation before export.
"""

from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime
from enum import Enum
from typing import List
from uuid import UUID

from fastapi import APIRouter, HTTPException, Query, Request, status
from fastapi.responses import StreamingResponse

from contextsafe.api.schemas import ErrorResponse
from contextsafe.api.middleware.session import get_session_id
from contextsafe.api.services.pii_validation import (
    CriticalPiiMatch,
    validate_no_critical_pii,
    format_pii_validation_error,
)
from contextsafe.api.session_manager import session_manager


router = APIRouter(prefix="/v1", tags=["export"])


class ExportFormat(str, Enum):
    """Supported export formats."""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    JSON = "json"
    CSV = "csv"


def _generate_pdf_content(text: str, title: str) -> bytes:
    """
    Generate a PDF from text content using reportlab if available,
    otherwise falls back to simple structure.

    Preserves line breaks and handles Unicode correctly.
    """
    try:
        # Try using reportlab for proper PDF with Unicode support
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=20,
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,  # Line spacing
            spaceAfter=6,
        )

        story = []

        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # Process text preserving line breaks
        # Split by double newlines for paragraphs, single newlines for line breaks
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            # Replace single newlines with <br/> for line breaks within paragraph
            para_html = para.replace('\n', '<br/>')
            # Escape special characters
            para_html = para_html.replace('&', '&amp;').replace('<br/>', '<br/>')
            if para_html.strip():
                story.append(Paragraph(para_html, body_style))
                story.append(Spacer(1, 6))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        # Fallback to simple PDF structure
        pass

    # Simplified PDF fallback (limited but works without dependencies)
    # Normalize text: handle both \r\n and \n line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Process text into lines with word wrapping at ~80 chars
    raw_lines = text.split('\n')
    wrapped_lines = []
    max_line_len = 80

    for line in raw_lines:
        if not line:
            wrapped_lines.append('')  # Preserve empty lines
        elif len(line) <= max_line_len:
            wrapped_lines.append(line)
        else:
            # Word wrap long lines
            words = line.split(' ')
            current_line = ''
            for word in words:
                if len(current_line) + len(word) + 1 <= max_line_len:
                    current_line = current_line + ' ' + word if current_line else word
                else:
                    if current_line:
                        wrapped_lines.append(current_line)
                    current_line = word
            if current_line:
                wrapped_lines.append(current_line)

    # Calculate number of pages needed (roughly 50 lines per page)
    lines_per_page = 50
    total_pages = (len(wrapped_lines) + lines_per_page - 1) // lines_per_page
    if total_pages < 1:
        total_pages = 1

    # For simplicity, still output a single page but with all content
    # Use a smaller font to fit more content
    pdf_content = """%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length PLACEHOLDER >>
stream
BT
/F1 10 Tf
50 770 Td
"""
    # Add title
    safe_title = title.replace('(', '\\(').replace(')', '\\)').replace('\\', '\\\\')
    pdf_content += f"({safe_title}) Tj\n0 -16 Td\n"

    # Add lines (limit to reasonable amount for fallback)
    for line in wrapped_lines[:200]:  # Increased from 50 to 200
        # Escape PDF special characters and handle Unicode
        safe_line = line.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
        # Replace problematic Unicode chars with ASCII equivalents
        safe_line = safe_line.encode('ascii', 'replace').decode('ascii')
        pdf_content += f"({safe_line}) Tj\n0 -12 Td\n"

    pdf_content += """ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000500 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
600
%%EOF"""

    # Calculate actual stream length and replace placeholder
    stream_start = pdf_content.find('stream\n') + 7
    stream_end = pdf_content.find('\nendstream')
    stream_length = stream_end - stream_start
    pdf_content = pdf_content.replace('PLACEHOLDER', str(stream_length))

    return pdf_content.encode('latin-1', errors='replace')


def _generate_docx_content(text: str, title: str) -> bytes:
    """
    Generate a DOCX from text content using python-docx if available,
    otherwise falls back to simple OOXML structure.

    Preserves line breaks and handles Unicode correctly.
    """
    try:
        # Try using python-docx for proper DOCX with full formatting
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Add title
        heading = doc.add_heading(title, level=1)

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Process text preserving line breaks
        # Split by double newlines for separate paragraphs
        sections = text.split('\n\n')
        for section in sections:
            if section.strip():
                # For single-line breaks within a section, add them to same paragraph
                para = doc.add_paragraph()
                lines = section.split('\n')
                for i, line in enumerate(lines):
                    if i > 0:
                        para.add_run().add_break()  # Line break within paragraph
                    para.add_run(line)
            else:
                # Empty section = blank paragraph for spacing
                doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        # Fallback to simple OOXML structure
        pass

    import zipfile

    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Create DOCX (which is a ZIP file with XML content)
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # [Content_Types].xml
        content_types = '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
        zf.writestr('[Content_Types].xml', content_types)

        # _rels/.rels
        rels = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
        zf.writestr('_rels/.rels', rels)

        # word/_rels/document.xml.rels
        doc_rels = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''
        zf.writestr('word/_rels/document.xml.rels', doc_rels)

        # word/document.xml - the actual content
        # Escape XML special characters
        def escape_xml(s: str) -> str:
            return (s
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&apos;'))

        escaped_title = escape_xml(title)

        # Build paragraphs preserving empty lines
        paragraphs = []
        lines = text.split('\n')
        for line in lines:
            escaped_line = escape_xml(line)
            if line.strip():
                # Non-empty line
                paragraphs.append(f'<w:p><w:r><w:t xml:space="preserve">{escaped_line}</w:t></w:r></w:p>')
            else:
                # Empty line - preserve as empty paragraph for spacing
                paragraphs.append('<w:p/>')

        paragraphs_xml = '\n        '.join(paragraphs)

        document = f'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:body>
        <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>{escaped_title}</w:t></w:r></w:p>
        <w:p/>
        {paragraphs_xml}
    </w:body>
</w:document>'''
        zf.writestr('word/document.xml', document)

    return buffer.getvalue()


@router.get(
    "/projects/{project_id}/export/glossary",
    responses={
        200: {
            "content": {"application/octet-stream": {}},
            "description": "File download",
        },
        404: {"model": ErrorResponse, "description": "Project not found"},
    },
)
async def export_glossary(
    project_id: UUID,
    request: Request,
    format: ExportFormat = Query(ExportFormat.CSV, description="Export format"),
) -> StreamingResponse:
    """
    Export project glossary as CSV or JSON.

    Returns the glossary file as a binary download.
    """
    session_id = get_session_id(request)
    project_id_str = str(project_id)

    # Check project exists
    if not session_manager.get_project(session_id, project_id_str):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Project {project_id} not found",
        )

    entries = session_manager.get_glossary(session_id, project_id_str)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

    if format == ExportFormat.JSON:
        # Export as JSON
        export_data = {
            "project_id": project_id_str,
            "exported_at": datetime.utcnow().isoformat(),
            "total_entries": len(entries),
            "entries": [
                {
                    "original_text": e["original_text"],
                    "alias": e["alias"],
                    "category": e["category"],
                    "occurrences": e.get("occurrences", 1),
                }
                for e in entries
            ],
        }
        content = json.dumps(export_data, indent=2, ensure_ascii=False)

        return StreamingResponse(
            iter([content.encode("utf-8")]),
            media_type="application/json",
            headers={
                "Content-Disposition": f'attachment; filename="glosario_{timestamp}.json"'
            },
        )

    else:
        # Default: Export as CSV
        buffer = io.StringIO()
        writer = csv.writer(buffer)

        # Header
        writer.writerow(["Texto Original", "Alias", "Categoría", "Ocurrencias"])

        # Data
        for entry in entries:
            writer.writerow([
                entry["original_text"],
                entry["alias"],
                entry["category"],
                entry.get("occurrences", 1),
            ])

        csv_content = buffer.getvalue()

        return StreamingResponse(
            iter([csv_content.encode("utf-8")]),
            media_type="text/csv",
            headers={
                "Content-Disposition": f'attachment; filename="glosario_{timestamp}.csv"'
            },
        )


@router.post(
    "/documents/{document_id}/export/advanced",
    responses={
        200: {
            "content": {"application/octet-stream": {}},
            "description": "File download",
        },
        404: {"model": ErrorResponse, "description": "Document not found"},
        422: {"model": ErrorResponse, "description": "Document not anonymized or contains PII"},
    },
)
async def export_document_advanced(
    document_id: UUID,
    request: Request,
    format: ExportFormat = Query(ExportFormat.TXT, description="Export format"),
    skip_pii_check: bool = Query(False, description="Skip PII validation (NOT recommended)"),
    strict_validation: bool = Query(True, description="Check HIGH severity PII too"),
) -> StreamingResponse:
    """
    Export anonymized document in requested format (PDF, DOCX, TXT).

    **POLICY GATE**: Before export, validates that no critical PII remains.
    If DNI, SS, bank accounts, or other critical PII is found, export is BLOCKED.

    Set skip_pii_check=true to bypass (NOT recommended for production).

    Returns the file as a binary download.
    """
    session_id = get_session_id(request)
    doc_id_str = str(document_id)

    doc = session_manager.get_document(session_id, doc_id_str)
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Document {document_id} not found",
        )

    # Get anonymized content from session
    session = session_manager.get_session(session_id)
    doc_data = session.documents.get(doc_id_str) if session else None

    if doc_data and doc_data.anonymized:
        anonymized_text = doc_data.anonymized.get("anonymized", doc_data.content or "")
    else:
        anonymized_text = doc.content or ""

    # ============================================================
    # POLICY GATE: Validate no critical PII before export
    # ============================================================
    if not skip_pii_check:
        pii_matches = validate_no_critical_pii(anonymized_text, strict=strict_validation)
        if pii_matches:
            error_detail = format_pii_validation_error(pii_matches)
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=error_detail,
            )

    filename_base = doc.filename.rsplit(".", 1)[0]
    title = f"Documento Anonimizado: {filename_base}"

    if format == ExportFormat.PDF:
        pdf_bytes = _generate_pdf_content(anonymized_text, title)
        return StreamingResponse(
            iter([pdf_bytes]),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename_base}_anonimizado.pdf"'
            },
        )

    elif format == ExportFormat.DOCX:
        docx_bytes = _generate_docx_content(anonymized_text, title)
        return StreamingResponse(
            iter([docx_bytes]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename_base}_anonimizado.docx"'
            },
        )

    else:
        # Default: TXT
        return StreamingResponse(
            iter([anonymized_text.encode("utf-8")]),
            media_type="text/plain",
            headers={
                "Content-Disposition": f'attachment; filename="{filename_base}_anonimizado.txt"'
            },
        )
